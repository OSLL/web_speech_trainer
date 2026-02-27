import argparse
import logging
import os
import sys
from contextlib import nullcontext

import nltk
from docx import Document
from generator import VkrQuestionGenerator
from validator import VkrQuestionValidator
from logging_utils import setup_logging, log_timed, suppress_console_logs


def load_vkr_text(path: str) -> str:
    logger = logging.getLogger(__name__)
    if not os.path.exists(path):
        logger.error("Файл не найден: %s", path)
        sys.exit(1)

    with log_timed(logger, "чтение DOCX", путь=path):
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)

    logger.info("DOCX обработан: символов=%d абзацев=%d", len(text), len(doc.paragraphs))
    return text


def main():
    setup_logging()
    logger = logging.getLogger(__name__)

    parser = argparse.ArgumentParser()
    parser.add_argument("vkr_path")
    parser.add_argument("--no-overflow-logs", action="store_true")
    args = parser.parse_args()

    logger.info("Запуск генерации: файл=%s", args.vkr_path)

    with log_timed(logger, "проверка NLTK"):
        try:
            nltk.data.find("tokenizers/punkt_tab/english")
        except LookupError:
            logger.info("Загрузка данных NLTK")
            nltk.download("punkt_tab")
            nltk.download("stopwords")

    text = load_vkr_text(args.vkr_path)

    with log_timed(logger, "инициализация генератора"):
        gen = VkrQuestionGenerator(args.vkr_path)

    with log_timed(logger, "инициализация валидатора"):
        validator = VkrQuestionValidator(text)

    with log_timed(logger, "генерация вопросов"):
        questions = gen.generate_all()

    logger.info("Сгенерировано вопросов: %d", len(questions))

    quiet = suppress_console_logs() if args.no_overflow_logs else nullcontext()

    with quiet:
        for idx, q in enumerate(questions, 1):
            if q.startswith("---"):
                print(f"\n{q}")
                continue

            rel = validator.check_relevance(q)
            clr = validator.check_clarity(q)
            diff = validator.check_difficulty(q)

            logger.info(
                "Вопрос %d релевантность=%s ясность=%s сложность=%s",
                idx,
                rel,
                clr,
                diff,
            )

            print(f"\n {q}")
            print(f"  - релевантность: {rel}")
            print(f"  - ясность:   {clr}")
            print(f"  - сложность:{diff}")


if __name__ == "__main__":
    main()
